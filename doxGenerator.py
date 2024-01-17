import datetime
from docx import Document
import subprocess


def generate_docx_with_bullets(question,answer,content,output_folder="",srcs=None):
    output_path=output_folder+(str(datetime.datetime.now())[:19].replace(" ","_").replace(":","_"))+".docx"
    doc = Document()
    
    doc.add_heading(question, level=2)
    
    # Split the content into lines
    lines = content.split('\n')
    p = doc.add_paragraph(answer)

    # Add paragraphs with bullets
    for line in lines:
        if line.strip():  # Check if the line is not empty
            p = doc.add_paragraph()
            p.add_run(u'\u2022').bold = True  # Bullet character (use '\u2022' for a solid bullet)
            p.add_run(f' {line}')
    
    if srcs is not None:
        doc.add_heading('Sources:', level=1)
        for srcs in srcs:
            doc.add_paragraph(f'• {srcs}')
        
    # Save the document
    doc.save(output_path)
    subprocess.Popen(['start', 'WINWORD.EXE', output_path], shell=True) 
    print(f"Document '{output_path}' generated successfully.")

# generate_docx_with_bullets(content="generate_legal_notice( response)",output_folder="./Outputs/")

if __name__ == "__main__":
    content_with_bullets = """
    This is a sample document with bullets.
    - Bullet point 1
    - Bullet point 2
    - Bullet point 3
    """

    generate_docx_with_bullets(content_with_bullets,heading="hello",output_folder="./tmp/",source_file="xyz.txt",page_no=25)
    
