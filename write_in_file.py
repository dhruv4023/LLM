import datetime
from docx import Document
import subprocess


def generate_docx_with_bullets(heading=None,main_paragraph=None,content=None,output_folder="",srcs=None):
    output_path=output_folder+(str(datetime.datetime.now())[:19].replace(" ","_").replace(":","_"))+".docx"
    doc = Document()
    
    if heading is not None: doc.add_heading(heading, level=2)

    if main_paragraph is not None: p = doc.add_paragraph(main_paragraph)

    if content is not None:
        # Add paragraphs with bullets
        for line in content.split('\n'):
            if line.strip():  # Check if the line is not empty
                p = doc.add_paragraph()
                p.add_run(u'\u2022').bold = True  # Bullet character (use '\u2022' for a solid bullet)
                p.add_run(f' {line}')
    
    if srcs is not None:
        doc.add_heading('Sources:', level=1)
        for srcs in srcs:
            doc.add_paragraph(f'• {srcs}')
        
    # Save the document
    doc.save(output_path)
    subprocess.Popen(['start', 'WINWORD.EXE', output_path], shell=True) 
    print(f"Document '{output_path}' generated successfully.")

# generate_docx_with_bullets(content="generate_legal_notice( response)",output_folder="./Outputs/")

# if __name__ == "__main__":
#     content_with_bullets = """
#     This is a sample document with bullets.
#     - Bullet point 1
#     - Bullet point 2
#     - Bullet point 3
#     """

#     generate_docx_with_bullets(content_with_bullets,heading="hello",output_folder="./tmp/",source_file="xyz.txt",page_no=25)
    

import subprocess

def open_in_notepad(text):
    # Write text to tmp.txt
    with open('tmp.txt', 'w') as file:
        file.write(text)

    # Open tmp.txt using Notepad
    subprocess.Popen(['notepad.exe', 'tmp.txt'])